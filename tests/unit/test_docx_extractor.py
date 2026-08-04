from pathlib import Path

import pytest

from folderscribe.domain.extraction import ExtractionStatus
from folderscribe.infrastructure.extractors.docx_extractor import DocxTextExtractor


@pytest.fixture
def extractor() -> DocxTextExtractor:
    return DocxTextExtractor()


class TestDocxExtractor:
    def test_supports_docx(self, extractor: DocxTextExtractor) -> None:
        assert extractor.supports(Path("file.docx"))
        assert extractor.supports(Path("file.DOCX"))
        assert not extractor.supports(Path("file.doc"))
        assert not extractor.supports(Path("file.pdf"))

    def test_name_and_version(self, extractor: DocxTextExtractor) -> None:
        assert extractor.name == "docx"
        assert extractor.version == "1"

    def test_extract_paragraphs(
        self, tmp_path: Path, extractor: DocxTextExtractor
    ) -> None:
        from docx import Document

        f = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(f))

        result = extractor.extract(f)
        assert result.status == ExtractionStatus.EXTRACTED
        assert "First paragraph" in (result.content or "")
        assert "Second paragraph" in (result.content or "")

    def test_extract_tables(self, tmp_path: Path, extractor: DocxTextExtractor) -> None:
        from docx import Document

        f = tmp_path / "table.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = "B1"
        table.rows[1].cells[0].text = "A2"
        table.rows[1].cells[1].text = "B2"
        doc.save(str(f))

        result = extractor.extract(f)
        assert result.status == ExtractionStatus.EXTRACTED
        assert "A1" in (result.content or "")
        assert "B2" in (result.content or "")

    def test_extract_empty(self, tmp_path: Path, extractor: DocxTextExtractor) -> None:
        from docx import Document

        f = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(f))

        result = extractor.extract(f)
        assert result.status == ExtractionStatus.EXTRACTED_EMPTY
        assert result.content == ""

    def test_extract_unicode(self, tmp_path: Path, extractor: DocxTextExtractor) -> None:
        from docx import Document

        f = tmp_path / "unicode.docx"
        doc = Document()
        doc.add_paragraph("Café ñoño 世界")
        doc.save(str(f))

        result = extractor.extract(f)
        assert result.status == ExtractionStatus.EXTRACTED
        assert "Café" in (result.content or "")
        assert "世界" in (result.content or "")

    def test_cancel(self, tmp_path: Path, extractor: DocxTextExtractor) -> None:
        from docx import Document

        f = tmp_path / "cancel.docx"
        doc = Document()
        doc.add_paragraph("test")
        doc.save(str(f))
        cancelled = False

        def cancel() -> bool:
            return cancelled

        cancelled = True
        result = extractor.extract(f, cancel_check=cancel)
        assert result.status == ExtractionStatus.CANCELLED

    def test_truncation(self, tmp_path: Path, extractor: DocxTextExtractor) -> None:
        from docx import Document

        f = tmp_path / "long.docx"
        doc = Document()
        doc.add_paragraph("A" * 500)
        doc.save(str(f))

        result = extractor.extract(f, max_chars=100)
        assert result.status == ExtractionStatus.PARTIAL
        assert result.is_truncated
