import logging
from collections.abc import Callable
from pathlib import Path

from folderscribe.domain.extraction import ExtractionStatus, TextExtraction
from folderscribe.domain.interfaces import TextExtractor

logger = logging.getLogger(__name__)


class DocxTextExtractor(TextExtractor):
    def __init__(self) -> None:
        self._name = "docx"
        self._version = "1"

    @property
    def name(self) -> str:
        return self._name

    @property
    def version(self) -> str:
        return self._version

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def extract(
        self,
        path: Path,
        max_chars: int = 100000,
        max_pages: int = 500,
        max_read_bytes: int = 10485760,
        cancel_check: Callable[[], bool] | None = None,
    ) -> TextExtraction:
        if cancel_check is not None and cancel_check():
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.CANCELLED,
            )

        try:
            from docx import Document
        except ImportError:
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.ERROR,
                error_code="missing_dependency",
                error_message="python-docx is not installed",
            )

        try:
            stat = path.stat()
            file_size = stat.st_size

            doc = Document(str(path))
        except Exception as e:
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.ERROR,
                error_code="docx_open_error",
                error_message=str(e),
            )

        if cancel_check is not None and cancel_check():
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.CANCELLED,
            )

        parts: list[str] = []
        char_count = 0
        is_truncated = False

        for para in doc.paragraphs:
            if cancel_check is not None and cancel_check():
                return TextExtraction(
                    absolute_path=path,
                    extractor=self._name,
                    extractor_version=self._version,
                    config_version="",
                    status=ExtractionStatus.CANCELLED,
                )
            text = para.text.strip()
            if text:
                parts.append(text)
                char_count += len(text)
                if char_count >= max_chars:
                    is_truncated = True
                    break

        if not is_truncated:
            for table in doc.tables:
                if cancel_check is not None and cancel_check():
                    return TextExtraction(
                        absolute_path=path,
                        extractor=self._name,
                        extractor_version=self._version,
                        config_version="",
                        status=ExtractionStatus.CANCELLED,
                    )
                for row in table.rows:
                    row_texts: list[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        line = " | ".join(row_texts)
                        parts.append(line)
                        char_count += len(line)
                        if char_count >= max_chars:
                            is_truncated = True
                            break
                if is_truncated:
                    break

        if cancel_check is not None and cancel_check():
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.CANCELLED,
            )

        content = "\n".join(parts)
        char_count = len(content)
        is_truncated = char_count > 0 and char_count >= max_chars

        if char_count == 0:
            return TextExtraction(
                absolute_path=path,
                extractor=self._name,
                extractor_version=self._version,
                config_version="",
                status=ExtractionStatus.EXTRACTED_EMPTY,
                content="",
                char_count=0,
                file_size=file_size,
            )

        status: ExtractionStatus = ExtractionStatus.EXTRACTED
        partial_reason: str | None = None
        is_partial = False
        if is_truncated:
            status = ExtractionStatus.PARTIAL
            is_partial = True
            partial_reason = f"reached_char_limit_{max_chars}"

        return TextExtraction(
            absolute_path=path,
            extractor=self._name,
            extractor_version=self._version,
            config_version="",
            status=status,
            content=content,
            char_count=char_count,
            file_size=file_size,
            is_truncated=is_truncated,
            is_partial=is_partial,
            partial_reason=partial_reason,
        )
